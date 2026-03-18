"""RAG: загрузка docs (PDF/DOCX), поиск по FAISS (OpenRouter) или по ключевым словам с fallback."""
import os
from typing import List, Dict, Any, Optional

from pypdf import PdfReader
import docx

# Пути и лимиты
DOCS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "docs"))
FAISS_INDEX_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), "rag_faiss_index"))
MAX_RAG_CONTEXT_CHARS = 6000
DEFAULT_TOP_K = 10
OPENROUTER_EMBEDDINGS_BASE = "https://openrouter.ai/api/v1"
EMBEDDING_MODEL = "openai/text-embedding-3-small"

# Заголовок, который подставляется перед выдержками из базы в контекст LLM
RAG_HEADER = (
    "Вот выдержки из локальной базы знаний (документы из папки docs). "
    "Отвечай, опираясь прежде всего на них. Если информации недостаточно, "
    "честно скажи, чего не хватает.\n\n"
)


def _split_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Режет текст на фрагменты заданного размера с перекрытием для RAG."""
    chunks: List[str] = []
    start = 0
    length = len(text)
    if length == 0:
        return chunks

    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        start += max(chunk_size - overlap, 1)
    return chunks


def _load_pdf(path: str) -> str:
    """Читает PDF-файл, возвращает объединённый текст страниц."""
    try:
        reader = PdfReader(path)
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages_text)
    except Exception as e:
        print(f"[RAG] Ошибка чтения PDF {path}: {e}")
        return ""


def _load_docx(path: str) -> str:
    """Читает DOCX-файл, возвращает текст параграфов."""
    try:
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)
    except Exception as e:
        print(f"[RAG] Ошибка чтения DOCX {path}: {e}")
        return ""


def _format_context_parts(items: List[tuple[str, str]], limit: int) -> List[str]:
    """Форматирует список (источник, текст) в блоки для контекста с лимитом по символам."""
    parts = []
    total = 0
    for idx, (source, text) in enumerate(items, start=1):
        part = f"[{idx}] ({source})\n{text}\n"
        if total + len(part) > limit:
            break
        parts.append(part)
        total += len(part)
    return parts


def build_knowledge_base(docs_dir: str | None = None) -> List[Dict[str, Any]]:
    """Обходит папку docs (PDF/DOCX), режет на чанки, возвращает список {text, source}."""
    base_dir = docs_dir or DOCS_DIR
    chunks: List[Dict[str, Any]] = []
    files_added: List[str] = []
    files_skipped: List[str] = []  # формат не поддерживается
    files_failed: List[str] = []   # ошибка чтения или пустой текст

    if not os.path.isdir(base_dir):
        print(f"[RAG] Папка с базой знаний не найдена: {base_dir}")
        return chunks

    for root, _, files in os.walk(base_dir):
        for name in files:
            lower = name.lower()
            path = os.path.join(root, name)

            text = ""
            if lower.endswith(".pdf"):
                text = _load_pdf(path)
            elif lower.endswith(".docx"):
                text = _load_docx(path)
            else:
                files_skipped.append(name)
                continue

            if not text:
                files_failed.append(name)
                continue

            files_added.append(name)
            for part in _split_text(text):
                chunks.append(
                    {
                        "text": part,
                        "source": name,
                    }
                )

    print(f"[RAG] Добавлено в базу знаний: {len(files_added)} файлов, {len(chunks)} фрагментов.")
    for name in files_added:
        print(f"  — {name}")
    if files_skipped:
        print(f"[RAG] Пропущено (формат не поддерживается .pdf/.docx): {len(files_skipped)} — {', '.join(files_skipped)}")
    if files_failed:
        print(f"[RAG] Не загружено (ошибка или пустой текст): {len(files_failed)} — {', '.join(files_failed)}")
    return chunks


def get_embeddings():
    """Создаёт клиент эмбеддингов OpenRouter (для FAISS). При отсутствии ключа — None."""
    try:
        from langchain_openai import OpenAIEmbeddings
        api_key = os.getenv("OPENROUTER_API_KEY")
        if not api_key:
            return None
        return OpenAIEmbeddings(
            openai_api_key=api_key,
            openai_api_base=OPENROUTER_EMBEDDINGS_BASE,
            model=EMBEDDING_MODEL,
        )
    except Exception as e:
        print(f"[RAG] Эмбеддинги недоступны: {e}")
        return None


def build_faiss_index(
    knowledge_chunks: List[Dict[str, Any]],
    index_path: str = FAISS_INDEX_PATH,
) -> Optional[Any]:
    """Строит FAISS-индекс по чанкам, сохраняет в index_path, возвращает vectorstore."""
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document

    embeddings = get_embeddings()
    if not embeddings or not knowledge_chunks:
        return None

    docs = [
        Document(page_content=c["text"], metadata={"source": c["source"]})
        for c in knowledge_chunks
    ]
    print("[RAG] Построение FAISS-индекса (эмбеддинги через OpenRouter)...")
    try:
        vectorstore = FAISS.from_documents(docs, embeddings)
        vectorstore.save_local(index_path)
        print(f"[RAG] FAISS-индекс сохранён: {index_path}")
        return vectorstore
    except Exception as e:
        print(f"[RAG] Ошибка построения FAISS: {e}")
        return None


def load_faiss_index(index_path: str = FAISS_INDEX_PATH) -> Optional[Any]:
    """Загружает ранее сохранённый FAISS-индекс с диска."""
    from langchain_community.vectorstores import FAISS

    if not os.path.isdir(index_path):
        return None
    embeddings = get_embeddings()
    if not embeddings:
        return None
    try:
        # Не используем dangerous deserialization: снижает риск,
        # если индекс повреждён/подменён.
        return FAISS.load_local(index_path, embeddings)
    except Exception as e:
        print(f"[RAG] Ошибка загрузки FAISS: {e}")
        return None


def load_or_build_faiss_index(knowledge_chunks: List[Dict[str, Any]]) -> Optional[Any]:
    """Загружает FAISS с диска или строит заново по чанкам."""
    store = load_faiss_index()
    if store is not None:
        print("[RAG] Используется FAISS-индекс (семантический поиск).")
        return store
    return build_faiss_index(knowledge_chunks)


def _normalize_word(w: str) -> str:
    """Приводит слово к нижнему регистру, оставляет буквы, цифры, дефис и ё."""
    return "".join(c for c in w.lower() if c.isalnum() or c in "ё-")


def _score_chunk(query_tokens: List[str], chunk_text: str) -> int:
    """Оценка релевантности чанка запросу: совпадение слов даёт баллы (fallback без FAISS)."""
    text_lower = chunk_text.lower()
    chunk_words = [_normalize_word(w) for w in text_lower.split() if len(w) > 1]
    score = 0
    for token in query_tokens:
        if not token:
            continue
        t = _normalize_word(token)
        if len(t) < 2:
            continue
        if t in text_lower:
            score += 2
        elif any(t in w or w in t for w in chunk_words):
            score += 1
    return score


def _fallback_chunks(chunks: List[Dict[str, Any]], n: int = 5) -> List[Dict[str, Any]]:
    """Выбирает до n чанков из разных источников, если по ключевым словам ничего не подошло."""
    seen_sources: set[str] = set()
    result: List[Dict[str, Any]] = []
    for c in chunks:
        if c["source"] not in seen_sources and len(result) < n:
            result.append(c)
            seen_sources.add(c["source"])
        if len(result) >= n:
            break
    if len(result) < n:
        for c in chunks:
            if c not in result and len(result) < n:
                result.append(c)
    return result


def retrieve_context(
    knowledge_chunks: List[Dict[str, Any]],
    query: str,
    k: int | None = None,
    max_chars: int | None = None,
    vectorstore: Optional[Any] = None,
) -> str:
    """Ищет по запросу: при наличии vectorstore — FAISS, иначе по ключевым словам + fallback. Возвращает текст для контекста LLM."""
    k = k if k is not None else DEFAULT_TOP_K
    limit = max_chars or MAX_RAG_CONTEXT_CHARS

    if vectorstore is not None:
        try:
            docs = vectorstore.similarity_search(query, k=k)
            items = [(doc.metadata.get("source", "?"), doc.page_content) for doc in docs]
            parts = _format_context_parts(items, limit)
            if parts:
                return RAG_HEADER + "\n\n".join(parts)
        except Exception as e:
            print(f"[RAG] Ошибка FAISS, fallback на ключевые слова: {e}")

    if not knowledge_chunks:
        return ""

    tokens = [w for t in query.split() if t.strip() for w in [_normalize_word(t)] if len(w) >= 2]
    scored = [(_score_chunk(tokens, c["text"]), c) for c in knowledge_chunks]
    scored.sort(key=lambda x: x[0], reverse=True)
    top_chunks = [c for _, c in scored[:k]] if scored and scored[0][0] > 0 else _fallback_chunks(knowledge_chunks, n=k)
    parts = _format_context_parts([(c["source"], c["text"]) for c in top_chunks], limit)
    if not parts:
        return ""
    return RAG_HEADER + "\n\n".join(parts)
