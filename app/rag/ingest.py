"""Загрузка документов (PDF/DOCX), чанкирование, построение FAISS-индекса."""
from __future__ import annotations

import os
from typing import Any, Dict, List, Optional

from pypdf import PdfReader
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS

from app.config import rag_settings
from app.rag.preprocess import preprocess_text
from app.rag.embeddings import get_embeddings


def _project_root() -> str:
    """Корень проекта (папка LLM_bot)."""
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))


def _get_docs_dir() -> str:
    if rag_settings.docs_dir:
        return os.path.abspath(rag_settings.docs_dir)
    return os.path.join(_project_root(), "docs")


def _get_faiss_index_path() -> str:
    if rag_settings.faiss_index_path:
        return os.path.abspath(rag_settings.faiss_index_path)
    return os.path.join(_project_root(), "rag_faiss_index")


def _load_pdf(path: str) -> str:
    """Читает PDF-файл, возвращает объединённый текст страниц."""
    try:
        reader = PdfReader(path)
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages_text)
    except Exception as e:
        print(f"[RAG] Ошибка чтения PDF {path}: {e}")
        return ""


def _load_docx(path: str) -> str:
    """Читает DOCX-файл, возвращает текст параграфов."""
    try:
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)
    except Exception as e:
        print(f"[RAG] Ошибка чтения DOCX {path}: {e}")
        return ""


def _create_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=rag_settings.chunk_size,
        chunk_overlap=rag_settings.chunk_overlap,
        length_function=len,
    )


def build_knowledge_base(docs_dir: str | None = None) -> List[Dict[str, Any]]:
    """
    Обходит папку docs (PDF/DOCX), предобрабатывает текст,
    разбивает RecursiveCharacterTextSplitter, возвращает список {text, source}.
    """
    base_dir = docs_dir or _get_docs_dir()
    chunks: List[Dict[str, Any]] = []
    files_added: List[str] = []
    files_skipped: List[str] = []
    files_failed: List[str] = []

    if not os.path.isdir(base_dir):
        print(f"[RAG] Папка с базой знаний не найдена: {base_dir}")
        return chunks

    splitter = _create_splitter()

    for root, _, files in os.walk(base_dir):
        for name in files:
            lower = name.lower()
            path = os.path.join(root, name)

            raw_text = ""
            if lower.endswith(".pdf"):
                raw_text = _load_pdf(path)
            elif lower.endswith(".docx"):
                raw_text = _load_docx(path)
            else:
                files_skipped.append(name)
                continue

            if not raw_text:
                files_failed.append(name)
                continue

            text = preprocess_text(raw_text)
            if not text:
                files_failed.append(name)
                continue

            files_added.append(name)
            split_chunks = splitter.split_text(text)
            for part in split_chunks:
                if part.strip():
                    chunks.append({"text": part.strip(), "source": name})

    print(f"[RAG] Добавлено в базу знаний: {len(files_added)} файлов, {len(chunks)} фрагментов.")
    for name in files_added:
        print(f"  — {name}")
    if files_skipped:
        print(f"[RAG] Пропущено (формат не поддерживается): {len(files_skipped)} — {', '.join(files_skipped)}")
    if files_failed:
        print(f"[RAG] Не загружено (ошибка или пустой текст): {len(files_failed)} — {', '.join(files_failed)}")
    return chunks


def build_faiss_index(
    knowledge_chunks: List[Dict[str, Any]],
    index_path: str | None = None,
) -> Optional[Any]:
    """Строит FAISS-индекс по чанкам, сохраняет в index_path, возвращает vectorstore."""
    embeddings = get_embeddings()
    if not embeddings or not knowledge_chunks:
        return None

    path = index_path or _get_faiss_index_path()
    docs = [
        Document(page_content=c["text"], metadata={"source": c["source"]})
        for c in knowledge_chunks
    ]
    print("[RAG] Построение FAISS-индекса (эмбеддинги через OpenRouter)...")
    try:
        vectorstore = FAISS.from_documents(docs, embeddings)
        vectorstore.save_local(path)
        print(f"[RAG] FAISS-индекс сохранён: {path}")
        return vectorstore
    except Exception as e:
        print(f"[RAG] Ошибка построения FAISS: {e}")
        return None


def load_faiss_index(index_path: str | None = None) -> Optional[Any]:
    """Загружает ранее сохранённый FAISS-индекс с диска."""
    path = index_path or _get_faiss_index_path()
    if not os.path.isdir(path):
        return None
    embeddings = get_embeddings()
    if not embeddings:
        return None
    try:
        return FAISS.load_local(
            path,
            embeddings,
            allow_dangerous_deserialization=True,
        )
    except Exception as e:
        print(f"[RAG] Ошибка загрузки FAISS: {e}")
        return None


def load_or_build_faiss_index(knowledge_chunks: List[Dict[str, Any]]) -> Optional[Any]:
    """Загружает FAISS с диска или строит заново по чанкам."""
    store = load_faiss_index()
    if store is not None:
        print("[RAG] Используется FAISS-индекс (семантический поиск).")
        return store
    return build_faiss_index(knowledge_chunks)
